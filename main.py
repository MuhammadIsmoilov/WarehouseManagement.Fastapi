from fastapi import FastAPI, HTTPException,  Response,Depends
from fastapi.middleware.cors import CORSMiddleware
from lib.auth import permission_dependency
from src.models import (RoleInsert, RoleUpdate, UserUpdate, 
                        UsersRegistration, InvoiceInsert, 
                        InsertCategory, UpdateCategory, 
                        InsertCustomer, InsertSupplier, UpdateSupplier, 
                        UpdateCustomer,GetTransactions,UpdateIncoming_and_outgoing,
                        InsertPermission, UpdatePermission, UserLogin,
                        RolesPermissionInsert,Incoming_and_outgoingInsert,
                        RolesPermissionUpdate,ProductInsert,
                        ProductUpdate,InsertNomenclature,UpdateNomenclature,
                        GetStock,PurchaseInsert,FilterProduct,
                        PurchasePropertyInsert,PurchasePropertyUpdate,
                        PurchaseUpdate,InvoicePropertyInsert,
                        InvoicePropertyUpdate,InvoiceUpdate,InsertInventory,UpdateInventory,Tags)
from lib.connection import connection
import lib.acl as ACL
from fastapi.responses import FileResponse
from docx import Document
import json

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=['*'],  
    allow_credentials=True,
    allow_methods=['*'],
    allow_headers=['*']
)

@app.post('/user/registration/', tags=[Tags.user])


@app.post('/login/', tags=[Tags.login])
async def login(user: UserLogin, response: Response ):
    try:
        with connection() as cur:
            cur.execute('CALL users.authentication(%s, %s, %s, %s)', (user.phone_number, user.password, 0, 0))
            result = cur.fetchone()
            returnResult = {
                'access_token': '',
                'refresh_token': ''
            }
        if result[0] == 0:
            returnResult['user_id'] = result[1]
            returnResult['access_token'] = ACL.access_token(user.phone_number, result[1])
            response.set_cookie(key='access_token', value=returnResult['access_token'])
            returnResult['refresh_token'] = ACL.refresh_token(user.phone_number, result[1])
            response.set_cookie(key='refresh_token', value=returnResult['refresh_token'])
        return returnResult
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    

@app.post('/user/reg/', tags=[Tags.login])
def user_registration(data: UsersRegistration):
    try:
        with connection() as cur:
            text_res = "Initial Value"
            cur.execute('CALL users.user_insert(%s, %s, %s, %s, %s, %s, %s)', 
                        (data.fullname, data.gender, data.user_address, 
                         data.user_phone_number, data.user_password,data.role_id, text_res))
            
            result = cur.fetchone()[0]
        
        return {"message": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    

@app.put('/update/user/', tags=[Tags.user])
def update_user(data: UserUpdate,
    permission_check: None = Depends(permission_dependency('AllowAll','manage_users','update_user'))  
    ):
    try:
        with connection() as cur:
            text_res = "Initial Value"
            cur.execute('call users.update_user(%s, %s, %s, %s, %s, %s, %s,)', (data.user_id,data.fullname, data.gender,data.user_address, 
                         data.user_phone_number, data.user_password,data.role_id, text_res,))
        return 'User updated successfully'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    








@app.post('/insert/role/',tags=[Tags.user])
def role_insert(
    data: RoleInsert,
    permission_check: None = Depends(permission_dependency('AllowAll','manage_roles','create_role'))
):
    try:
        with connection() as cur:
            cur.execute('CALL users.role_insert (%s)', (data.role_name,))
            return 'Role added successfully'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))





@app.put('/update/role/',tags=[Tags.user])
def update_role(
    data: RoleUpdate,
    permission_check: None = Depends(permission_dependency('manage_roles','update_role','AllowAll'))
    
    ):
    try:
        with connection() as cur:
            cur.execute('CALL users.update_role(%s, %s)', (data.role_id, data.role_name))
        return 'Role updated successfully'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    
@app.post('/insert/permission',tags=[Tags.user])
def insert_permission(
    data:InsertPermission,
    permission_check: None = Depends(permission_dependency('manage_permission','create_permission','AllowAll')),  
    ):
    try:
        with connection() as curr:
            curr.execute('call users.permission_insert(%s)',(data.permission_name,))
            return 'Ok'
    except Exception as e :
       raise HTTPException(status_code=500,detail=str(e))     
    
   


@app.post('/update/permission',tags=[Tags.user])
def update_permission(
    data:UpdatePermission,
    permission_check: None = Depends(permission_dependency('manage_permission','update_permission','AllowAll')),  
    ):
    try:
        with connection() as curr:
            curr.execute('call users.update_permission(%s,%s)',(data.permission_id,data.permission_name,))
            return 'Ok'
    except Exception as e :
       raise HTTPException(status_code=500,detail=str(e)) 




@app.post('/insert/role-permission/',tags=[Tags.user])
def roles_permission_insert(
    data:RolesPermissionInsert,
    permission_check: None = Depends(permission_dependency('manage_role_permission','create_role_permission','AllowAll')) 

    ):
    try:
        with connection() as cur:
            cur.execute('call users.roles_permission_insert(%s,%s)', (data.role_id,data.permission_id,))
            return 'Role permission added seccessfuly'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



@app.put('/update/role-permission/',tags=[Tags.user])
def update_roles_permission(
    data: RolesPermissionUpdate,
    permission_check: None = Depends(permission_dependency('manage_role_permission','update_role_permission','AllowAll'))
    ):
    try:
        with connection() as cur:
            cur.execute('call users.update_roles_permission(%s,%s,%s)', (data.roles_permission_id,data.role_id, data.permission_id,))
        return 'Role permission updated successfully'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    



@app.post('/insert/invoice/',tags=[Tags.sales])
def insert_invoice(
    data: InvoiceInsert,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_invoice','create_invoice','AllowAll'))
    ):
    try:
        with connection() as conn:
            with conn as cur:
                cur.execute(
                    'CALL sales.invoice_insert(%s, %s, %s, %s)', 
                    (data.customer_id,  data.invoice_quantity, data.invoice_unit_price, data.invoice_date,))
        return {"message": 'Invoice added successfully'}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
    
@app.put('/update/invoice/',tags=[Tags.sales])
def update_invoice(
    data: InvoiceUpdate,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_invoice','update_invoice','AllowAll'))
    ):
    try:
        with connection() as cur:
         cur.execute('CALL sales.update_invoice(%s, %s, %s, %s, %s, %s)', 
                       (data.invoice_id,
                        data.customer_id,
                        data.invoice_quantity, 
                        data.invoice_unit_price,
                        data.invoice_date,))
        return 'Invoice updated successfully'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))




@app.post('/insert/invoice-property/',tags=[Tags.sales])
def insert_invoice_property(
    data: InvoicePropertyInsert,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_invoice','insert-invoice-property','AllowAll'))
    ):
    try:
        with connection() as conn:
            with conn as cur:
                cur.execute(
                    'CALL sales.invoice_property_insert(%s, %s)', 
                    (data.prod_id, data.invoice_id,))
        return {"message": 'Invoice property added successfully'}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
    

@app.put('/update/invoice-property/',tags=[Tags.sales])
def update_invoice_property(
    data: InvoicePropertyUpdate,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_invoice','update-invoice-property','AllowAll'))
    ):
    try:
        with connection() as cur:
         cur.execute('CALL sales.update_invoice_property(%s, %s, %s)', 
                       (data.property_id,
                        data.prod_id,
                        data.invoice_id,))
        return 'Invoice property updated successfully'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



@app.get('/invoice/of/product/',tags=[Tags.sales])
def get_invoice_details_of_product(
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_invoice','invoice-product','AllowAll'))
):
    result = None
    with connection() as cur:
        cur.callproc('sales.get_invoice_details_of_product')
        result = cur.fetchall()
        
        if not result or not result[0]:
            raise HTTPException(status_code=404, detail="No data found")

        row = result[0][0]

        invoice = row[0]

        doc = Document("files/Накладная продажи.docx")
        data_mapping = {
            "{invoice_id}": str(invoice['invoice_id']),
            "{invoice_date}": str(invoice['invoice_date']),
            "{custom_name}": str(invoice['custom_name']),
            "{custom_address}": str(invoice['custom_address']),
            "{custom_contact_info}": str(invoice['custom_contact_info']),
            "{invoice_quantity}": str(invoice['invoice_quantity']),
            "{invoice_unit_price}": str(invoice['invoice_unit_price']),
            "{articul}": str(invoice['articul']),
            "{barcode}": str(invoice['barcode']),
            "{ctg_name}": str(invoice['ctg_name']),
            "{invoice_total_amount}": str(invoice['invoice_total_amount'])
        }

        column_to_property = {
            0: "nom_name",
            1: "invoice_quantity",
            2: "invoice_unit_price",
            3: "articul",
            4: "barcode",
            5: "invoice_total_amount"
        }

        tbl = doc.add_table(0, 6)

        for r in row:
            t_row = tbl.add_row()
            for y in range(0, 6):
                t_row.cells[y].add_paragraph(str(r[column_to_property[y]]))

        # Replace placeholders in DOCX
        for paragraph in doc.paragraphs:
            for marker, value in data_mapping.items():
                if marker in paragraph.text:
                    paragraph.text = paragraph.text.replace(marker, value)

        filename_docx = f"{invoice['invoice_id']}.docx"
        doc.save(filename_docx)
        doc = Document(filename_docx)

        return FileResponse(filename_docx, media_type='application/pdf', filename=filename_docx)




@app.post('/insert/purchase/',tags=[Tags.purchase])
def insert_purchase(
    data: PurchaseInsert,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_purchase','insert_purchase','AllowAll'))
    ):
    try:
        with connection() as cur:
           cur.execute('CALL buying.purchase_insert(%s, %s, %s, %s)', 
                            (data.sup_id,
                             data.purch_date, 
                             data.purch_quantity,
                             data.purch_unit_price))
        return {"message": 'Purchase added successfuly'}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
    

@app.put('/update/purchase/',tags=[Tags.purchase])
def update_purchase(
    data: PurchaseUpdate,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_purchase','update_purchase','AllowAll'))
    ):
    try:
        with connection() as cur:
         cur.execute('CALL buying.update_purchase(%s, %s, %s, %s, %s, %s)', 
                       (data.purch_id,
                        data.sup_id,
                        data.purch_date, 
                        data.purch_quantity,
                        data.purch_unit_price,
                        data.user_id,))
        return 'Purchase updated successfully'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))





@app.post('/insert/purchase-property/',tags=[Tags.purchase])
def insert_purchase_property(
    data: PurchasePropertyInsert,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_purchase','insert_purchase-property','AllowAll'))
    ):
    try:
        with connection() as cur:
           cur.execute('CALL buying.purchase_property_insert(%s, %s)', 
                            (data.prod_id,
                             data.purchase_id,))
        return {"message": 'Purchase property added successfuly'}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
    

@app.put('/update/purchase-property/',tags=[Tags.purchase])
def update_purchase_property(
    data: PurchasePropertyUpdate,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_purchase','update_purchase-property','AllowAll'))
    ):
    try:
        with connection() as cur:
         cur.execute('CALL buying.update_purchase_property(%s, %s, %s)', 
                       (data.property_id,
                        data.prod_id,
                        data.purchase_id,))
        return 'Purchase property updated successfully'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))




@app.get('/purchase-invoice/of/product/',tags=[Tags.purchase])
def get_purchase_details_of_product(
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_purchase','invoice-purchase-product','AllowAll'))
):
    result = None
    with connection() as cur:
        cur.callproc('buying.get_purchase_details_of_product')
        result = cur.fetchall()
        
        if not result or not result[0]:
            raise HTTPException(status_code=404, detail="No data found")



        row = result[0][0]
        purchase = row[0]
        
        doc = Document("files/Накладная покупки.docx")
        data_mapping = {
            "{purchase_id}":str(purchase['purchase_id']),
            "{purchase_date}": str(purchase['purchase_date']),
            "{supp_name}": str(purchase['supp_name']),
            "{supp_address}": str(purchase['supp_address']),
            "{supp_contact_info}": str(purchase['supp_contact_info']),
            "{nom_name}": str(purchase['nom_name']),
            "{purchase_quantity}": str(purchase['purchase_quantity']),
            "{purchase_unit_price}": str(purchase['purchase_unit_price']),
            "{articul}": str(purchase['articul']),
            "{barcode}": str(purchase['barcode']),
            "{ctg_name}": str(purchase['ctg_name']),
            "{purchase_total_amount}": str(purchase['purchase_total_amount']),
            "{fullname}":str(purchase['fullname'])
        }

        column_to_property = {
            0: "nom_name",
            1: "purchase_quantity",
            2: "purchase_unit_price",
            3: "articul",
            4: "barcode",
            5: "purchase_total_amount"
            
        }

        tbl = doc.add_table(0, 6)

        for r in row:
            t_row = tbl.add_row()
            for y in range(0, 6):
                t_row.cells[y].add_paragraph(str(r[column_to_property[y]]))
                
        # Replace placeholders in DOCX
        for paragraph in doc.paragraphs:
            for marker, value in data_mapping.items():
                if marker in paragraph.text:
                    paragraph.text = paragraph.text.replace(marker, value)

        filename_docx = f"{purchase['purchase_id']}.docx"
        doc.save(filename_docx)
        doc = Document(filename_docx)

        return FileResponse(filename_docx, media_type='application/pdf', filename=filename_docx)


    
@app.post('/insert/customer',tags=[Tags.customer])   
def insert_customer(
    data:InsertCustomer,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_customer','create_customer','AllowAll'))  
    ):
    try:
        with connection() as cur:
            cur.execute('call customer.customer_insert(%s,%s,%s)',(data.custom_name,data.custom_address,data.custom_custom_contact_info,))
            return 'Ok'
    except Exception as e:
        raise HTTPException(status_code=500,detail=str(e))    


@app.put('/update/customer',tags=[Tags.customer])   
def update_customer(
    data:UpdateCustomer,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_customer','update_customer','AllowAll'))
    ):
    try:
        with connection() as cur:
            cur.execute('call customer.update_customer(%s,%s,%s,%s)',(data.custom_id,data.custom_name,data.custom_addres,data.custom_custom_contact_info,))
            return 'Ok'
    except Exception as e:
        raise HTTPException(status_code=500,detail=str(e))  





@app.post('/insert/supplier',tags=[Tags.provider])   
def insert_supplier(
    data:InsertSupplier,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_supplier','create_supplier','AllowAll')) 
    ):
    try:
        with connection() as cur:
            cur.execute('call provider.supplier_insert(%s,%s,%s)',(data.supp_name,data.supp_address,data.supp_contact_info,))
            return 'Ok'
    except Exception as e:
        raise HTTPException(status_code=500,detail=str(e))    
        



@app.put('/update/supplier',tags=[Tags.provider])   
def update_supplier(
    data:UpdateSupplier,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_supplier','update_supplier','AllowAll'))  
    ):
    try:
        with connection() as cur:
            cur.execute('call provider.update_supplier(%s,%s,%s,%s)',(data .supp_id,data.supp_name,data.supp_address,data.supp_contact_info,))
            return 'Ok'
    except Exception as e:
        raise HTTPException(status_code=500,detail=str(e))  

 





@app.post('/insert/product/',tags=[Tags.catalog])
def product_insert(data: ProductInsert,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_product','create_product','AllowAll'))
                   
    ):
    try:
        with connection() as cur:
            cur.execute('CALL catalogs.product_insert(%s,%s,%s,%s,%s,%s,%s)', 
                        (data.nom_id, data.category_id, data.prod_price, 
                         data.prod_quontity, data.articul, data.barcode, 
                         json.dumps(data.prod_attribute)))  
            return 'Product added successfully'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))    



@app.put('/update/product/',tags=[Tags.catalog])
def update_product(data: ProductUpdate,
                   token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_product','update_product','AllowAll'))
                   ):
    try:
        with connection() as cur:
            cur.execute('call catalogs.update_product(%s,%s,%s,%s,%s,%s,%s,%s)',
                         (data.prod_id,data.nom_id, data.category_id,data.prod_price,
                          data.prod_quontity,data.articul,data.barcode,
                          json.dumps(data.prod_attribute)))
        return 'Product updated successfully'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))





@app.post('/insert/nomenclature/',tags=[Tags.catalog])
def nomenclature_insert(
    data: InsertNomenclature,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_nomenclature','create_nomenclature','AllowAll')) 
    ):
    try:
        with connection() as cur:
            cur.execute('call catalogs.nomenclature_insert(%s,%s)', (data.nom_name, data.nom_description))
            return 'Nomenclature added successfully'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    



@app.put('/update/nomenclature/',tags=[Tags.catalog])
def update_nomenclature(
    data: UpdateNomenclature,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_nomenclature','update_nomenclature','AllowAll'))
    ):
    try:
        with connection() as cur:
            cur.execute('call catalogs.update_nomenclature(%s,%s, %s)', (data.nom_id,data.nom_name, data.nom_description,))
        return 'Nomenclature updated successfully'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))





@app.post('/insert/category/',tags=[Tags.catalog])
def category_insert(
    data: InsertCategory,
    token: str = Depends(ACL.JWTBearer()),
    permission_check_1: None = Depends(permission_dependency('manage_category','create_category','AllowAll'))
    ):
    try:
        with connection() as cur:
            cur.execute('call catalogs.category_insert(%s,%s)', (data.ctg_name, data.parent_category_id))
            return 'Category added successfully'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    


@app.put('/update/category/',tags=[Tags.catalog])
def update_category(
    data: UpdateCategory,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_category','update_category','AllowAll'))
    ):
    try:
        with connection() as cur:
            cur.execute('call catalogs.update_category(%s,%s, %s)', (data.ctg_id,data.ctg_name, data.parent_category_id,))
        return 'Category updated successfully'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))








@app.post('/get/stock/',tags=[Tags.inventory])
def get_stock(data:GetStock):
    try:
        with connection() as conn:
            with conn as cur:
                cur.execute('SELECT * FROM inventory.get_stock(%s,%s)',
                            (data.nom_name,data.prod_value,))
                results = cur.fetchone()[0]
        return results
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    


@app.post('/search/product/',tags=[Tags.catalog])
def search_product(data:FilterProduct):
    try:
        with connection() as conn:
            with conn as cur:
                cur.execute('SELECT * FROM catalogs.filter_product(%s,%s,%s,%s,%s,%s)',
                            (data.nom_name,data.ctg_name,data.prod_price,data.articul,data.barcode,data.prod_properties,))
                results = cur.fetchone()[0]
        return results
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post('/get/transaction/',tags=[Tags.inventory])
def get_transaction(data:GetTransactions):
    try:
        with connection() as conn:
            with conn as cur:
                cur.execute('SELECT * FROM inventory.get_transaction(%s,%s,%s)',
                            (data.nom_name,data.type_of_income_or_outgoing,data.status,))
                results = cur.fetchone()[0]
        return results
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



@app.post('/income-outgoing/', tags=[Tags.inventory])
def insert_income_and_outgoing(data: Incoming_and_outgoingInsert):
    try:
        with connection() as cur:
           cur.execute('CALL inventory.incoming_and_outgoing_insert(%s, %s, %s, %s, %s)', 
                           (data.product_id,
                            data.quantity,
                            data.date_of_incoming_and_outgoing,
                            data.incoming_and_outgoing_type,
                            data.is_active,))
        return {"message": 'Income and outgoing added successfuly'}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
    
@app.put('/income-outgoing/', tags=[Tags.inventory])
def update_income_and_outgoing(data: UpdateIncoming_and_outgoing):
    try:
        with connection() as cur:
         cur.execute('CALL inventory.update_incoming_and_outgoing(%s, %s, %s, %s, %s, %s, %s)', 
                           (data.income_id,
                            data.product_id,
                            data.quantity,
                            data.date_of_incoming_and_outgoing,
                            data.incoming_and_outgoing_type,
                            data.is_active,))
        return 'Income and outgoing updated successfully'
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))







@app.post('/insert/inventory',tags=[Tags.inventory])
def insert_incentory(
    data:InsertInventory,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_inventory','insert_inventory','AllowAll')) 
    ):


    try:
        with connection() as conn:
            with conn as cur:
                cur.execute('call inventory.inventory_insert(%s,%s,%s)',
                            (data.product_id,data.inven_quantity,data.updated_at,))
        return "Inventory added successfully"
    except Exception as e:
        raise HTTPException(status_code=500,detail=str(e))

@app.put('/update/inventory',tags=[Tags.inventory])
def update_inventory(
    data:UpdateInventory,
    token: str = Depends(ACL.JWTBearer()),
    permission_check: None = Depends(permission_dependency('manage_inventory','update_inventory','AllowAll')) 
                     ):
    try:
        with connection() as conn:
            with conn as cur:
                cur.execute('call inventory.update_inventory(%s,%s,%s,%s,%s)',
                            (data.inven_id,data.product_id,data.transaction_type,data.transaction_date,)) 
                results = cur.fetchone() [0]
        return results 
    except Exception as e:
        raise HTTPException(status_code=500,detail=str(e))      

















@app.get('/categories/all',tags=[Tags.catalog])
def get_all_categories(
    token: str = Depends(ACL.JWTBearer()),
    permission_check_1: None = Depends(permission_dependency('AllowAll','get_categories')),  

):
    result = None 
    with connection() as cur:
        cur.execute('SELECT catalogs.get_all_categories()')
        result = cur.fetchone()[0]

    return result



@app.get('/permission/all',tags=[Tags.user])
def get_all_permission(
    token: str = Depends(ACL.JWTBearer()),
    permission_check_1: None = Depends(permission_dependency('AllowAll','get_permissions')),  
):
    result = None 
    with connection() as cur:
        cur.execute('SELECT users.get_all_permission()')
        result = cur.fetchone()[0]

    return result

@app.get('/role/all',tags=[Tags.user])
def get_all_role(
    token: str = Depends(ACL.JWTBearer()),
    permission_check_1: None = Depends(permission_dependency('AllowAll','get_roles')),  
):
    result = None 
    with connection() as cur:
        cur.execute('SELECT users.get_all_role()')
        result = cur.fetchone()[0]

    return result


@app.get('/users/all',tags=[Tags.user])
def get_all_users(
    token: str = Depends(ACL.JWTBearer()),
    permission_check_1: None = Depends(permission_dependency('AllowAll','get_users')),  
):
    result = None 
    with connection() as cur:
        cur.execute('SELECT users.get_all_users()')
        result = cur.fetchone()[0]

    return result


